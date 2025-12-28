from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt


# Set column widths to closely match the PDF layout
# 1st column: serial number (narrow)
# 2nd column: labels
# 3rd column: values (widest)
def set_col_widths(table):
    widths = [Inches(0.45), Inches(2.6), Inches(3.45)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


# Remove extra spacing inside table cells
# This helps achieve a compact, legal-form style layout
def compact_cell(cell):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1


# Center content vertically inside a table cell
def center_vertical(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def create_doc(output_path):
    # Create a new Word document
    doc = Document()

    # Adjust page margins so the table spans the page like the PDF
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Helper function to add centered, bold headings
    def center_bold(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

    # Main heading section
    center_bold("FORM ‘A’")
    center_bold("MEDIATION APPLICATION FORM")

    p = doc.add_paragraph("[REFER RULE 3(1)]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("Mumbai District Legal Services Authority")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("City Civil Court, Mumbai")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Small gap before the table starts
    doc.add_paragraph("")

    # Create the main table structure
    table = doc.add_table(rows=18, cols=3)
    table.style = "Table Grid"
    set_col_widths(table)

    # Reduce font size inside the table for a dense legal-form look
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # ---- DETAILS OF PARTIES ----
    table.cell(0, 0).merge(table.cell(0, 2))
    table.cell(0, 0).text = "DETAILS OF PARTIES:"
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    compact_cell(table.cell(0, 0))

    # Applicant name row
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "Name of Applicant"
    table.cell(1, 2).text = "{{client_name}}"

    for c in range(3):
        compact_cell(table.cell(1, c))
        center_vertical(table.cell(1, c))

    # Applicant address header
    table.cell(2, 1).merge(table.cell(2, 2))
    table.cell(2, 1).text = "Address and contact details of Applicant"
    table.cell(2, 1).paragraphs[0].runs[0].bold = True
    compact_cell(table.cell(2, 1))

    # Registered address
    table.cell(3, 1).text = "REGISTERED ADDRESS:"
    table.cell(3, 1).paragraphs[0].runs[0].bold = True
    table.cell(3, 2).text = "{{branch_address}}"
    compact_cell(table.cell(3, 1))
    compact_cell(table.cell(3, 2))

    # Correspondence address
    table.cell(4, 1).text = "CORRESPONDENCE BRANCH ADDRESS:"
    table.cell(4, 1).paragraphs[0].runs[0].bold = True
    table.cell(4, 2).text = "{{branch_address}}"
    compact_cell(table.cell(4, 1))
    compact_cell(table.cell(4, 2))

    # Contact details
    table.cell(5, 1).text = "Telephone No."
    table.cell(5, 2).text = "{{mobile}}"
    compact_cell(table.cell(5, 1))
    compact_cell(table.cell(5, 2))
    center_vertical(table.cell(5, 1))
    center_vertical(table.cell(5, 2))

    table.cell(6, 1).text = "Mobile No."
    table.cell(6, 2).text = ""
    compact_cell(table.cell(6, 1))
    compact_cell(table.cell(6, 2))

    table.cell(7, 1).text = "Email ID"
    table.cell(7, 2).text = "info@kslegal.co.in"
    compact_cell(table.cell(7, 1))
    compact_cell(table.cell(7, 2))

    # ---- OPPOSITE PARTY DETAILS ----
    table.cell(8, 0).merge(table.cell(8, 2))
    table.cell(8, 0).text = "Name, Address and Contact details of Opposite Party:"
    table.cell(8, 0).paragraphs[0].runs[0].bold = True
    compact_cell(table.cell(8, 0))

    table.cell(9, 1).merge(table.cell(9, 2))
    table.cell(9, 1).text = "Address and contact details of Defendant/s"
    table.cell(9, 1).paragraphs[0].runs[0].bold = True
    compact_cell(table.cell(9, 1))

    table.cell(10, 1).text = "Name"
    table.cell(10, 2).text = "{{customer_name}}"
    compact_cell(table.cell(10, 1))
    compact_cell(table.cell(10, 2))

    table.cell(11, 1).text = "REGISTERED ADDRESS:"
    table.cell(11, 1).paragraphs[0].runs[0].bold = True
    table.cell(11, 2).text = "{{address1 or '______________'}}"
    compact_cell(table.cell(11, 1))
    compact_cell(table.cell(11, 2))

    table.cell(12, 1).text = "CORRESPONDENCE ADDRESS:"
    table.cell(12, 1).paragraphs[0].runs[0].bold = True
    table.cell(12, 2).text = "{{address1 or '______________'}}"
    compact_cell(table.cell(12, 1))
    compact_cell(table.cell(12, 2))

    table.cell(13, 1).text = "Telephone No."
    table.cell(13, 2).text = ""
    compact_cell(table.cell(13, 1))
    compact_cell(table.cell(13, 2))

    table.cell(14, 1).text = "Mobile No."
    table.cell(14, 2).text = ""
    compact_cell(table.cell(14, 1))
    compact_cell(table.cell(14, 2))

    table.cell(15, 1).text = "Email ID"
    table.cell(15, 2).text = ""
    compact_cell(table.cell(15, 1))
    compact_cell(table.cell(15, 2))

    # ---- DETAILS OF DISPUTE ----
    table.cell(16, 0).merge(table.cell(16, 2))
    table.cell(16, 0).text = "DETAILS OF DISPUTE:"
    table.cell(16, 0).paragraphs[0].runs[0].bold = True
    compact_cell(table.cell(16, 0))

    table.cell(17, 0).merge(table.cell(17, 2))
    table.cell(17, 0).text = (
        "THE COMM. COURTS (PRE-INSTITUTION SETTLEMENT) RULES, 2018\n"
        "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"
    )
    compact_cell(table.cell(17, 0))

    # Save the generated Word document
    doc.save(output_path)


if __name__ == "__main__":
    create_doc("test_output.docx")
